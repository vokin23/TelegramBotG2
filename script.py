import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_channel_id(telegram_link):
  """Извлекает ID канала Telegram из ссылки."""

  # Проверка на наличие ссылки
  if not telegram_link:
    return None

  # Поиск ID в ссылке
  match = re.search(r't\.me\/([^\/]+)', telegram_link) or re.search(r'telegram\.me\/([^\/]+)', telegram_link)

  if match:
    return match.group(1)
  else:
    return None


def get_files_in_directory(directory):
  """Возвращает список путей файлов в директории."""
  return [os.path.join(directory, f) for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]


def create_docx_from_screenshots(screenshots, filename="screenshots.docx"):
  """
  Создает docx-файл с изображениями из списка скриншотов.

  Args:
    screenshots: Список путей к скриншотам.
    filename: Имя выходного docx-файла (по умолчанию "screenshots.docx").
  """

  # Создаем новый документ
  doc = Document()

  # Настройка шрифта для заголовка
  doc.styles['Normal'].font.name = 'Arial'
  doc.styles['Normal'].font.size = Pt(10)

  # Делим изображения на страницы по 25 изображений
  pages = [screenshots[i:i + 25] for i in range(0, len(screenshots), 25)]

  for page in pages:
    # Добавляем заголовок
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('Скриншоты')
    run.bold = True
    run.font.size = Pt(14)
    paragraph = doc.add_paragraph()

    # Добавляем изображения в один параграф
    for i, screenshot in enumerate(page):
      paragraph.add_run().add_picture(screenshot, width=Inches(1.8))
    doc.save(filename)

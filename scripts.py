# scripts.py
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import aiofiles


async def async_read_json_file(filename):
    async with aiofiles.open(filename, mode='r', encoding="utf-8") as file:
        data = await file.read()
        json_data = json.loads(data)
        return json_data


async def async_dump_json_file(filename, data):
    async with aiofiles.open(filename, mode='w', encoding='utf-8') as file:
        await file.write(json.dumps(data, ensure_ascii=False, indent=4))


def get_channel_id(telegram_link):
    """Извлекает ID канала Telegram из ссылки."""
    if not telegram_link:
        return None
    match = re.search(r't\.me\/([^\/]+)', telegram_link) or re.search(r'telegram\.me\/([^\/]+)', telegram_link)
    if match:
        return match.group(1)
    else:
        return None


def get_files_in_directory(directory):
    """Возвращает список путей файлов в директории."""
    if not os.path.exists(directory):
        os.makedirs(directory)
    return [os.path.join(directory, f) for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]


def create_docx_from_screenshots(screenshots, filename="screenshots.docx"):
    """Создает docx-файл с изображениями из списка скриншотов."""
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)
    pages = [screenshots[i:i + 25] for i in range(0, len(screenshots), 25)]
    for page in pages:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('Скриншоты')
        run.bold = True
        run.font.size = Pt(14)
        paragraph = doc.add_paragraph()
        for i, screenshot in enumerate(page):
            paragraph.add_run().add_picture(screenshot, width=Inches(1.8))
        doc.save(filename)